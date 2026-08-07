from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import Response
from sqlalchemy.orm import Session
from pydantic import BaseModel
from typing import Optional
from datetime import datetime
import json
import io

from models.database import get_db, Script, NewsSource, User
from routers.auth import get_current_user
from services.gemini_agent import regenerate_script_with_angle

router = APIRouter(prefix="/api/scripts", tags=["scripts"])


# ── Schemas ────────────────────────────────────────────────────────────────────
class ScriptCreate(BaseModel):
    news_source_id: int
    platform: str
    angle: Optional[str] = None
    angle_reasoning: Optional[str] = None
    headline: Optional[str] = None
    content: str
    tone: Optional[str] = None
    word_count: Optional[int] = None


class ScriptUpdate(BaseModel):
    headline: Optional[str] = None
    content: Optional[str] = None
    angle: Optional[str] = None


class ScriptResponse(BaseModel):
    id: int
    news_source_id: int
    user_id: int
    platform: str
    angle: Optional[str]
    angle_reasoning: Optional[str]
    headline: Optional[str]
    content: str
    tone: Optional[str]
    word_count: Optional[int]
    version: int
    created_at: datetime
    updated_at: datetime

    class Config:
        from_attributes = True


class RegenerateRequest(BaseModel):
    news_source_id: int
    angle_index: int  # 0, 1, or 2 (top3 angles)
    platform: str


# ── Routes ─────────────────────────────────────────────────────────────────────
@router.get("/", response_model=list[ScriptResponse])
def list_scripts(
    platform: Optional[str] = None,
    skip: int = 0,
    limit: int = 50,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Daftar semua naskah tim (semua user bisa lihat)."""
    query = db.query(Script).order_by(Script.updated_at.desc())
    if platform:
        query = query.filter(Script.platform == platform)
    return query.offset(skip).limit(limit).all()


@router.post("/", response_model=ScriptResponse)
def create_script(
    script_data: ScriptCreate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Simpan naskah baru."""
    # Verify news source exists
    news = db.query(NewsSource).filter(NewsSource.id == script_data.news_source_id).first()
    if not news:
        raise HTTPException(status_code=404, detail="Berita sumber tidak ditemukan")

    script = Script(
        news_source_id=script_data.news_source_id,
        user_id=current_user.id,
        platform=script_data.platform,
        angle=script_data.angle,
        angle_reasoning=script_data.angle_reasoning,
        headline=script_data.headline,
        content=script_data.content,
        tone=script_data.tone,
        word_count=script_data.word_count or len(script_data.content.split())
    )
    db.add(script)
    db.commit()
    db.refresh(script)
    return script


@router.get("/{script_id}", response_model=ScriptResponse)
def get_script(
    script_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    script = db.query(Script).filter(Script.id == script_id).first()
    if not script:
        raise HTTPException(status_code=404, detail="Naskah tidak ditemukan")
    return script


@router.put("/{script_id}", response_model=ScriptResponse)
def update_script(
    script_id: int,
    update_data: ScriptUpdate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Edit naskah — siapapun anggota tim bisa edit."""
    script = db.query(Script).filter(Script.id == script_id).first()
    if not script:
        raise HTTPException(status_code=404, detail="Naskah tidak ditemukan")

    if update_data.headline is not None:
        script.headline = update_data.headline
    if update_data.content is not None:
        script.content = update_data.content
        script.word_count = len(update_data.content.split())
    if update_data.angle is not None:
        script.angle = update_data.angle

    script.version += 1
    script.updated_at = datetime.utcnow()
    db.commit()
    db.refresh(script)
    return script


@router.delete("/{script_id}")
def delete_script(
    script_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    script = db.query(Script).filter(Script.id == script_id).first()
    if not script:
        raise HTTPException(status_code=404, detail="Naskah tidak ditemukan")

    if script.user_id != current_user.id and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Tidak punya izin menghapus naskah ini")

    db.delete(script)
    db.commit()
    return {"message": "Naskah berhasil dihapus"}


@router.post("/regenerate")
async def regenerate_script(
    request: RegenerateRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Regenerate naskah dengan angle berbeda yang dipilih user."""
    news = db.query(NewsSource).filter(NewsSource.id == request.news_source_id).first()
    if not news:
        raise HTTPException(status_code=404, detail="Berita sumber tidak ditemukan")

    ai_analysis = json.loads(news.ai_analysis)
    top3_angles = ai_analysis.get("top3_angles", [])

    if request.angle_index >= len(top3_angles):
        raise HTTPException(status_code=400, detail="Index angle tidak valid")

    selected_angle = top3_angles[request.angle_index]
    facts_data = json.loads(news.extracted_facts)

    try:
        script_data = await regenerate_script_with_angle(facts_data, selected_angle, request.platform)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI Error: {str(e)}")

    return script_data


@router.get("/{script_id}/export")
def export_script(
    script_id: int,
    format: str = "txt",
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Export naskah ke TXT atau DOCX."""
    script = db.query(Script).filter(Script.id == script_id).first()
    if not script:
        raise HTTPException(status_code=404, detail="Naskah tidak ditemukan")

    platform_names = {
        "tv_radio": "TV-Radio",
        "article": "Artikel",
        "instagram": "Instagram",
        "tiktok": "TikTok",
        "youtube": "YouTube-Shorts"
    }
    platform_label = platform_names.get(script.platform, script.platform)
    filename_base = f"naskah_{platform_label}_{script.id}"

    if format == "txt":
        content_txt = f"""NEWSSCRIPT AI — NASKAH BERITA
{'='*50}
Platform   : {platform_label}
Headline   : {script.headline or '-'}
Angle      : {script.angle or '-'}
Tone       : {script.tone or '-'}
Kata       : {script.word_count or 0}
Versi      : {script.version}
Dibuat     : {script.created_at.strftime('%d/%m/%Y %H:%M')}
Diperbarui : {script.updated_at.strftime('%d/%m/%Y %H:%M')}
{'='*50}

{script.content}
"""
        return Response(
            content=content_txt.encode("utf-8"),
            media_type="text/plain",
            headers={"Content-Disposition": f'attachment; filename="{filename_base}.txt"'}
        )

    elif format == "docx":
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()
            
            # Title
            title_para = doc.add_heading(f"Naskah Berita — {platform_label}", level=1)
            
            # Metadata table
            table = doc.add_table(rows=6, cols=2)
            table.style = "Table Grid"
            cells = [
                ("Platform", platform_label),
                ("Headline", script.headline or "-"),
                ("Angle", script.angle or "-"),
                ("Tone", script.tone or "-"),
                ("Jumlah Kata", str(script.word_count or 0)),
                ("Versi", str(script.version)),
            ]
            for i, (key, val) in enumerate(cells):
                table.cell(i, 0).text = key
                table.cell(i, 1).text = val

            doc.add_paragraph()
            doc.add_heading("Naskah:", level=2)
            doc.add_paragraph(script.content)

            # Save to bytes
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            return Response(
                content=buffer.read(),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f'attachment; filename="{filename_base}.docx"'}
            )
        except ImportError:
            raise HTTPException(status_code=500, detail="python-docx tidak terinstall")

    else:
        raise HTTPException(status_code=400, detail="Format tidak valid. Gunakan 'txt' atau 'docx'")
